"""
FastAPI Resume Parser Service
This service receives resume files (PDF/DOCX) as base64,
parses them using Groq LLM, and returns structured JSON.
No database operations - just parsing and returning data.
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import os
import base64
import tempfile
import pdfplumber
from docx import Document as DocxDocument
import requests
import json
import re
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="Resume Parser API", version="1.0.0")

# CORS middleware to allow .NET backend to call this service
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your .NET backend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# ===================== Models =====================

class EducationItem(BaseModel):
    degree: Optional[str] = ""
    institution: Optional[str] = ""
    year: Optional[str] = ""
    details: Optional[str] = ""

class ProjectItem(BaseModel):
    title: Optional[str] = ""
    description: Optional[str] = ""

class ExperienceItem(BaseModel):
    title: Optional[str] = ""
    company: Optional[str] = ""
    duration: Optional[str] = ""
    description: Optional[str] = ""

class ParsedResumeData(BaseModel):
    name: Optional[str] = ""
    email: Optional[str] = ""
    phone: Optional[str] = ""
    summary: Optional[str] = ""
    education: List[EducationItem] = []
    skills: List[str] = []
    projects: List[ProjectItem] = []
    experience: List[ExperienceItem] = []

class ResumeUploadRequest(BaseModel):
    filename: str
    file_base64: str

# ===================== Helper Functions =====================

def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes"""
    resume_text = ''
    tmp_path = None
    
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        print(f"[DEBUG] Created temporary PDF file: {tmp_path}")
        
        with pdfplumber.open(tmp_path) as pdf:
            print(f"[DEBUG] PDF has {len(pdf.pages)} pages")
            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        resume_text += text + '\n'
                        print(f"[DEBUG] Extracted {len(text)} characters from page {i+1}")
                except Exception as e:
                    print(f"[WARNING] Failed to extract text from page {i+1}: {str(e)}")
                    continue
    except Exception as e:
        print(f"[ERROR] PDF extraction failed: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
                print(f"[DEBUG] Cleaned up temporary file: {tmp_path}")
            except:
                pass
    
    return resume_text.strip()

def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes"""
    resume_text = ''
    tmp_path = None
    
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        print(f"[DEBUG] Created temporary DOCX file: {tmp_path}")
        
        doc = DocxDocument(tmp_path)
        print(f"[DEBUG] DOCX has {len(doc.paragraphs)} paragraphs")
        
        for para in doc.paragraphs:
            if para.text:
                resume_text += para.text + '\n'
        
        print(f"[DEBUG] Extracted {len(resume_text)} characters from DOCX")
    except Exception as e:
        print(f"[ERROR] DOCX extraction failed: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
                print(f"[DEBUG] Cleaned up temporary file: {tmp_path}")
            except:
                pass
    
    return resume_text.strip()

def extract_text_from_resume(filename: str, file_bytes: bytes) -> str:
    """Extract text based on file extension"""
    file_type = filename.split('.')[-1].lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf_bytes(file_bytes)
    elif file_type == 'docx':
        return extract_text_from_docx_bytes(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def extract_json_from_groq_response(response_text: str) -> dict:
    """Extract JSON from Groq response (handles markdown code blocks and text before JSON)"""
    if not response_text:
        print("[WARNING] Empty response text from Groq")
        return {}
    
    text = response_text.strip()
    original_text = text
    
    # Strategy 1: Look for JSON in markdown code blocks first (most reliable)
    # Handle both complete and incomplete code blocks
    if "```json" in text or "```" in text:
        # Find all code block markers
        code_blocks = list(re.finditer(r'```(?:json)?', text))
        if code_blocks:
            # Extract content from first code block
            start_marker = code_blocks[0].end()
            # Find the next ``` after start_marker, or use end of text if not found
            end_marker = text.find("```", start_marker)
            if end_marker == -1:
                # No closing ```, use end of text (incomplete JSON from Groq)
                end_marker = len(text)
                print("[DEBUG] No closing ``` found, using end of text (JSON might be incomplete)")
            
            text = text[start_marker:end_marker].strip()
            print("[DEBUG] Extracted JSON from markdown code block")
    
    text = text.strip()
    
    # Strategy 3: Find JSON object boundaries by matching braces (handles text before JSON)
    if not text.startswith("{"):
        start_idx = text.find("{")
        if start_idx != -1:
            # Count braces to find the matching closing brace
            brace_count = 0
            end_idx = -1
            for i in range(start_idx, len(text)):
                if text[i] == '{':
                    brace_count += 1
                elif text[i] == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        end_idx = i
                        break
            
            if end_idx != -1 and end_idx > start_idx:
                text = text[start_idx:end_idx+1]
                print("[DEBUG] Extracted JSON object from text with prefix (handled nested braces)")
    
    text = text.strip()
    
    # Strategy 4: Try to parse as-is
    try:
        parsed = json.loads(text)
        print(f"[DEBUG] Successfully parsed JSON with {len(parsed)} top-level keys")
        return parsed
    except json.JSONDecodeError as e:
        print(f"[WARNING] Initial JSON parse failed: {str(e)} at position {e.pos}")
        print(f"[DEBUG] Text around error (chars {max(0, e.pos-100)} to {min(len(text), e.pos+100)}):")
        print(f"  ...{text[max(0, e.pos-100):e.pos+100]}...")
        
        # Strategy 5: Fix common JSON issues
        try:
            fixed_text = text
            # Remove trailing commas
            fixed_text = re.sub(r',(\s*[}\]])', r'\1', fixed_text)
            # Fix unclosed strings (if any)
            # Remove any text after the last }
            last_brace = fixed_text.rfind('}')
            if last_brace != -1:
                fixed_text = fixed_text[:last_brace+1]
            
            parsed = json.loads(fixed_text)
            print("[DEBUG] Successfully parsed after fixing JSON issues (trailing commas, etc.)")
            return parsed
        except json.JSONDecodeError as e2:
            print(f"[ERROR] JSON decode error after fixes: {str(e2)} at position {e2.pos}")
            
            # Strategy 6: Try progressive truncation from the end
            # Sometimes Groq returns incomplete JSON or extra text
            brace_start = text.find('{')
            if brace_start != -1:
                print(f"[DEBUG] Attempting progressive truncation from position {len(text)}")
                # Try different end positions, starting from the last } and working backwards
                last_brace = text.rfind('}')
                if last_brace != -1:
                    # Try parsing from first { to last }
                    try:
                        potential_json = text[brace_start:last_brace+1]
                        parsed = json.loads(potential_json)
                        print(f"[DEBUG] Successfully parsed by using first {{ to last }}")
                        return parsed
                    except:
                        pass
                
                # Try truncating character by character from the end
                for truncate_pos in range(len(text), brace_start + 50, -1):
                    try:
                        potential_json = text[brace_start:truncate_pos]
                        # Check if it's valid JSON structure
                        if potential_json.count('{') == potential_json.count('}'):
                            parsed = json.loads(potential_json)
                            print(f"[DEBUG] Successfully parsed by truncating to position {truncate_pos}")
                            return parsed
                    except:
                        continue
            
            # Strategy 7: Try to fix incomplete JSON by closing arrays/objects intelligently
            try:
                fixed_text = text
                
                # Remove trailing commas first
                fixed_text = re.sub(r',(\s*[}\]])', r'\1', fixed_text)
                
                # Count unclosed brackets/braces
                open_braces = fixed_text.count('{') - fixed_text.count('}')
                open_brackets = fixed_text.count('[') - fixed_text.count(']')
                
                # If we have unclosed structures, try to fix them
                if open_braces > 0 or open_brackets > 0:
                    # Find where the JSON is incomplete and close it properly
                    # Close brackets first (they're inside objects), then braces
                    if open_brackets > 0:
                        # Find the last unclosed array
                        last_open_bracket = fixed_text.rfind('[')
                        if last_open_bracket != -1:
                            # Check if it has content
                            after_bracket = fixed_text[last_open_bracket+1:].strip()
                            if after_bracket and not after_bracket.startswith(']'):
                                # There's content, need to close the array
                                # Find a safe place to insert ]
                                # Look for the last complete item before the truncation
                                fixed_text = fixed_text.rstrip()
                                if not fixed_text.endswith(']'):
                                    fixed_text += ']'
                                    open_brackets -= 1
                    
                    # Close braces
                    if open_braces > 0:
                        fixed_text = fixed_text.rstrip()
                        if not fixed_text.endswith('}'):
                            # Find where to close - before the last incomplete structure
                            fixed_text += '\n' + '}' * open_braces
                    
                    # Close remaining brackets
                    if open_brackets > 0:
                        fixed_text += '\n' + ']' * open_brackets
                
                # Remove any trailing commas again after closing
                fixed_text = re.sub(r',(\s*[}\]])', r'\1', fixed_text)
                
                parsed = json.loads(fixed_text)
                print(f"[DEBUG] Successfully parsed after closing {open_braces} braces and {open_brackets} brackets")
                return parsed
            except Exception as e3:
                print(f"[WARNING] Failed to fix incomplete JSON: {str(e3)}")
                import traceback
                print(f"[DEBUG] Traceback: {traceback.format_exc()}")
            
            # Strategy 8: Last resort - try to extract partial data and return what we can
            print(f"[ERROR] All JSON extraction strategies failed")
            print(f"[DEBUG] Full response text length: {len(original_text)}")
            print(f"[DEBUG] First 1000 chars of original: {original_text[:1000]}")
            print(f"[DEBUG] Last 500 chars of original: {original_text[-500:]}")
            print(f"[DEBUG] Extracted text length: {len(text)}")
            print(f"[DEBUG] First 1000 chars of extracted: {text[:1000]}")
            
            # Return empty dict so caller can handle gracefully
            return {}

def call_groq_llm(resume_text: str) -> dict:
    """Send text to Groq LLM for JSON parsing"""
    if not GROQ_API_KEY:
        print("[ERROR] GROQ_API_KEY not set in environment")
        raise ValueError("GROQ_API_KEY not set in environment")
    
    if not resume_text or not resume_text.strip():
        print("[ERROR] Resume text is empty")
        raise ValueError("Resume text is empty")
    
    print(f"[DEBUG] Sending {len(resume_text)} characters to Groq API")
    
    prompt = f"""Parse the following resume text and return ONLY a complete, valid JSON object. Do not add any explanatory text, comments, or markdown formatting.

Required JSON structure:
{{
  "name": "",
  "email": "",
  "phone": "",
  "summary": "",
  "education": [
    {{
      "degree": "",
      "institution": "",
      "year": "",
      "details": ""
    }}
  ],
  "skills": [],
  "projects": [
    {{
      "title": "",
      "description": ""
    }}
  ],
  "experience": [
    {{
      "title": "",
      "company": "",
      "duration": "",
      "description": ""
    }}
  ]
}}

CRITICAL REQUIREMENTS:
- Return ONLY the JSON object, no other text
- Ensure the JSON is complete and valid
- All arrays and objects must be properly closed
- All strings must be properly quoted
- For skills: return a flat array of strings like ["Python", "Java"], not objects

Resume text to parse:
{resume_text}

Return the JSON now:"""
    
    headers = {
        'Authorization': f'Bearer {GROQ_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    data = {
        'messages': [
            {'role': 'system', 'content': 'You are a helpful assistant. Always return complete, valid JSON. Never truncate your response.'},
            {'role': 'user', 'content': prompt}
        ],
        'model': 'llama-3.1-8b-instant',
        'max_tokens': 4096  # Increased to prevent truncation
    }
    
    try:
        print("[DEBUG] Making request to Groq API...")
        response = requests.post(
            'https://api.groq.com/openai/v1/chat/completions',
            headers=headers,
            json=data,
            timeout=60
        )
        print(f"[DEBUG] Groq API response status: {response.status_code}")
        response.raise_for_status()
        
        response_data = response.json()
        if 'choices' not in response_data or not response_data['choices']:
            print("[ERROR] Groq response missing 'choices'")
            raise ValueError("Invalid response from Groq API: missing choices")
        
        content = response_data['choices'][0]['message']['content']
        print(f"[DEBUG] Groq returned {len(content)} characters of content")
        
        # Check if response was truncated (Groq stops at max_tokens)
        if 'finish_reason' in response_data['choices'][0]:
            finish_reason = response_data['choices'][0]['finish_reason']
            if finish_reason == 'length':
                print("[WARNING] Groq response was truncated (hit max_tokens limit). JSON may be incomplete.")
        
        parsed_json = extract_json_from_groq_response(content)
        
        if not parsed_json:
            print("[ERROR] Failed to extract JSON from Groq response")
            print(f"[DEBUG] Raw Groq content (first 1000 chars): {content[:1000]}")
            print(f"[DEBUG] Raw Groq content (last 500 chars): {content[-500:]}")
            raise ValueError("Groq returned empty or invalid JSON. The response may have been truncated.")
        
        print(f"[DEBUG] Successfully parsed JSON with {len(parsed_json)} keys")
        return parsed_json
    except requests.exceptions.RequestException as e:
        print(f"[ERROR] Groq API request failed: {str(e)}")
        raise Exception(f"Groq API request error: {str(e)}")
    except KeyError as e:
        print(f"[ERROR] Groq API response missing key: {str(e)}")
        raise Exception(f"Groq API response format error: {str(e)}")
    except Exception as e:
        print(f"[ERROR] Groq API error: {str(e)}")
        raise Exception(f"Groq API error: {str(e)}")

def generate_summary_with_groq(parsed_json: dict) -> str:
    """Generate a professional summary using Groq"""
    if not GROQ_API_KEY:
        return "Experienced professional with relevant technical expertise."
    
    prompt = f"""
Generate a concise professional summary for a person based on the
following structured resume data:

{json.dumps(parsed_json, indent=2)}

Guidelines:
- Do NOT include the person's name or contact info.
- Use a professional, factual tone.
- Mention the main role/specialization.
- If years of experience are clearly derivable from the data, include them;
  otherwise omit them (do not invent).
- Highlight key technical skills and what they enable the person to build.
- Keep the summary between 30 and 50 words.
- Output only the summary text.
"""
    
    headers = {
        'Authorization': f'Bearer {GROQ_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    data = {
        'messages': [
            {'role': 'system', 'content': 'You are a helpful assistant.'},
            {'role': 'user', 'content': prompt}
        ],
        'model': 'llama-3.1-8b-instant',
        'max_tokens': 150
    }
    
    try:
        response = requests.post(
            'https://api.groq.com/openai/v1/chat/completions',
            headers=headers,
            json=data,
            timeout=60
        )
        response.raise_for_status()
        content = response.json()['choices'][0]['message']['content']
        return content.strip()
    except:
        return "Experienced professional with relevant technical expertise."

# ===================== API Endpoints =====================

@app.get("/")
def read_root():
    """Health check endpoint"""
    return {
        "service": "Resume Parser API",
        "status": "running",
        "version": "1.0.0"
    }

@app.post("/parse-resume", response_model=ParsedResumeData)
async def parse_resume(request: ResumeUploadRequest):
    """
    Parse a resume file and return structured JSON data.
    
    Args:
        request: Contains filename and base64-encoded file content
    
    Returns:
        ParsedResumeData: Structured resume data
    """
    try:
        print(f"[DEBUG] Received request to parse resume: {request.filename}")
        
        # Validate inputs
        if not request.filename:
            raise HTTPException(status_code=400, detail="Filename is required")
        if not request.file_base64:
            raise HTTPException(status_code=400, detail="File base64 content is required")
        
        # Decode base64 file
        try:
            print(f"[DEBUG] Decoding base64 file (length: {len(request.file_base64)})")
            file_bytes = base64.b64decode(request.file_base64)
            print(f"[DEBUG] Decoded file size: {len(file_bytes)} bytes")
        except Exception as e:
            print(f"[ERROR] Base64 decode failed: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Invalid base64 encoding: {str(e)}")
        
        # Extract text from resume
        try:
            print(f"[DEBUG] Extracting text from {request.filename}")
            resume_text = extract_text_from_resume(request.filename, file_bytes)
            print(f"[DEBUG] Extracted text length: {len(resume_text)} characters")
        except Exception as e:
            print(f"[ERROR] Text extraction failed: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from file: {str(e)}")
        
        if not resume_text or not resume_text.strip():
            print("[ERROR] No text extracted from resume")
            raise HTTPException(status_code=400, detail="No text could be extracted from the resume")
        
        # Parse with Groq LLM
        try:
            print("[DEBUG] Calling Groq LLM to parse resume")
            parsed_json = call_groq_llm(resume_text)
            print(f"[DEBUG] Groq returned JSON with keys: {list(parsed_json.keys())}")
        except Exception as e:
            print(f"[ERROR] Groq LLM call failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to parse resume with AI: {str(e)}")
        
        if not parsed_json:
            print("[ERROR] Groq returned empty JSON")
            raise HTTPException(status_code=500, detail="AI parsing returned empty results")
        
        # Generate summary if not present or empty
        if not parsed_json.get("summary") or not str(parsed_json.get("summary", "")).strip():
            try:
                print("[DEBUG] Generating summary with Groq")
                summary = generate_summary_with_groq(parsed_json)
                parsed_json["summary"] = summary if summary else "Experienced professional with relevant technical expertise."
            except Exception as e:
                print(f"[WARNING] Summary generation failed: {str(e)}, using default")
                parsed_json["summary"] = "Experienced professional with relevant technical expertise."
        
        # Ensure all required fields exist with defaults
        parsed_json.setdefault("name", "")
        parsed_json.setdefault("email", "")
        parsed_json.setdefault("phone", "")
        parsed_json.setdefault("summary", "")
        parsed_json.setdefault("education", [])
        parsed_json.setdefault("skills", [])
        parsed_json.setdefault("projects", [])
        parsed_json.setdefault("experience", [])
        
        # Transform skills if they come as objects/dictionaries from Groq
        # Groq sometimes returns: [{"category": "...", "language": "...", "details": ["skill1", "skill2"]}]
        # We need: ["skill1", "skill2", "skill3", ...] - flat list of strings
        if "skills" in parsed_json and parsed_json["skills"]:
            original_skills_count = len(parsed_json["skills"])
            flattened_skills = []
            for skill_item in parsed_json["skills"]:
                if isinstance(skill_item, str):
                    # Already a string, keep it
                    flattened_skills.append(skill_item.strip())
                elif isinstance(skill_item, dict):
                    # Extract from "details" array if present (this is the main source of skills)
                    if "details" in skill_item and isinstance(skill_item["details"], list):
                        flattened_skills.extend([str(s).strip() for s in skill_item["details"] if s and str(s).strip()])
                    # Also extract from "language" field if it's a skill name (not a category)
                    if "language" in skill_item and skill_item["language"]:
                        lang = str(skill_item["language"]).strip()
                        # Skip generic category names
                        generic_categories = ["programming languages", "web & frameworks", "databases", "tools & platforms", "networking", "other skills"]
                        if lang.lower() not in generic_categories:
                            flattened_skills.append(lang)
            # Remove duplicates while preserving order, and filter out empty strings
            seen = set()
            unique_skills = []
            for skill in flattened_skills:
                if skill:
                    skill_lower = skill.lower().strip()
                    if skill_lower and skill_lower not in seen:
                        seen.add(skill_lower)
                        unique_skills.append(skill.strip())
            parsed_json["skills"] = unique_skills
            print(f"[DEBUG] Flattened skills: {original_skills_count} skill groups → {len(parsed_json['skills'])} individual skills")
        
        # Remove extra fields that Groq might return but we don't need (LinkedIn, GitHub, etc.)
        # Keep only the fields we expect
        allowed_fields = ["name", "email", "phone", "summary", "education", "skills", "projects", "experience"]
        parsed_json = {k: v for k, v in parsed_json.items() if k in allowed_fields}
        
        # Ensure lists contain proper objects
        if "education" in parsed_json and parsed_json["education"]:
            parsed_json["education"] = [
                {**item, "degree": item.get("degree", "") or "", "institution": item.get("institution", "") or "", 
                 "year": item.get("year", "") or "", "details": item.get("details", "") or ""}
                if isinstance(item, dict) else {"degree": "", "institution": "", "year": "", "details": ""}
                for item in parsed_json["education"]
            ]
        
        if "projects" in parsed_json and parsed_json["projects"]:
            parsed_json["projects"] = [
                {**item, "title": item.get("title", "") or "", "description": item.get("description", "") or ""}
                if isinstance(item, dict) else {"title": "", "description": ""}
                for item in parsed_json["projects"]
            ]
        
        if "experience" in parsed_json and parsed_json["experience"]:
            parsed_json["experience"] = [
                {**item, "title": item.get("title", "") or "", "company": item.get("company", "") or "",
                 "duration": item.get("duration", "") or "", "description": item.get("description", "") or ""}
                if isinstance(item, dict) else {"title": "", "company": "", "duration": "", "description": ""}
                for item in parsed_json["experience"]
            ]
        
        # Convert to Pydantic model for validation
        try:
            print("[DEBUG] Converting to Pydantic model")
            result = ParsedResumeData(**parsed_json)
            print("[DEBUG] Successfully created Pydantic model")
            return result
        except Exception as e:
            print(f"[ERROR] Pydantic validation failed: {str(e)}")
            print(f"[DEBUG] Parsed JSON that failed validation: {json.dumps(parsed_json, indent=2)}")
            raise HTTPException(status_code=500, detail=f"Data validation error: {str(e)}")
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except ValueError as ve:
        print(f"[ERROR] ValueError: {str(ve)}")
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"[ERROR] Unexpected error: {str(e)}")
        print(f"[ERROR] Traceback: {error_trace}")
        raise HTTPException(status_code=500, detail=f"Error parsing resume: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
